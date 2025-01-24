import gradio as gr
from transformers import AutoTokenizer, AutoModel
from openai import OpenAI
import os
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import tempfile
from astroquery.nasa_ads import ADS
import pyvo as vo
import pandas as pd

# Load the NASA-specific bi-encoder model and tokenizer
bi_encoder_model_name = "nasa-impact/nasa-smd-ibm-st-v2"
bi_tokenizer = AutoTokenizer.from_pretrained(bi_encoder_model_name)
bi_model = AutoModel.from_pretrained(bi_encoder_model_name)

# Set up OpenAI client
api_key = os.getenv('OPENAI_API_KEY')
client = OpenAI(api_key=api_key)

# Set up NASA ADS token
ADS.TOKEN = os.getenv('ADS_API_KEY')  # Ensure your ADS API key is stored in environment variables

# Define system message with instructions
system_message = """
You are ExosAI, a helpful assistant specializing in Exoplanet and Astrophysics research.
Generate a detailed structured response based on the following science context and user input, including the necessary observables, physical parameters, and technical requirements for observations. The response should include the following sections:
    Science Objectives: Describe key scientific study objectives related to the science context and user input.
    Physical Parameters: Outline the physical parameters related to the science context and user input.
    Observables: Specify the observables related to the science context and user input.
    Description of Desired Observations: Detail the types of observations related to the science context and user input.
    Technical Requirements Table: Generate a table with the following columns:
    - Requirements: The specific observational requirements (e.g., UV observations, Optical observations or Infrared observations).
    - Necessary: The necessary values or parameters (e.g., wavelength ranges, spatial resolution).
    - Desired: The desired values or parameters.
    - Justification: A scientific explanation of why these requirements are important.
    - Comments: Additional notes or remarks regarding each requirement.
    Example:
    | Requirements                     | Necessary                                | Desired                                  | Justification                                                                                                                                              | Comments                                                                                                         |
    |----------------------------------|------------------------------------------|------------------------------------------|------------------------------------------------------------------------------------------------------------------------------------------------------------|-----------------------------------------------------------------------------------------------------------------|
    | UV Observations                  | Wavelength: 1200–2100 Å, 2500–3300 Å     | Wavelength: 1200–3300 Å                  | Characterization of atomic and molecular emissions (H, C, O, S, etc.) from fluorescence and dissociative electron impact                                    | Needed for detecting H2O, CO, CO2, and other volatile molecules relevant for volatile delivery studies.         |
    | Infrared Observations            | Wavelength: 2.5–4.8 μm                   | Wavelength: 1.5–4.8 μm                   | Tracks water emissions and CO2 lines in icy bodies and small planetesimals                                                                                  | Also allows detection of 3 μm absorption feature in icy bodies.                                                |
    Ensure the response is structured clearly and the technical requirements table follows this format.
"""

def encode_text(text):
    inputs = bi_tokenizer(text, return_tensors='pt', padding=True, truncation=True, max_length=128)
    outputs = bi_model(**inputs)
    return outputs.last_hidden_state.mean(dim=1).detach().numpy().flatten()

def get_chunks(text, chunk_size=300):
    """
    Split a long piece of text into smaller chunks of approximately 'chunk_size' characters.
    """
    if not text.strip():
        raise ValueError("The provided context is empty or blank.")
    
    # Split the text into chunks of approximately 'chunk_size' characters
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    return chunks

def retrieve_relevant_context(user_input, context_texts, chunk_size=300, similarity_threshold=0.3):
    """
    Split the context text into smaller chunks, find the most relevant chunk
    using cosine similarity, and return the most relevant chunk.
    If no chunk meets the similarity threshold, return a fallback message.
    """
    # Check if the context is empty or just whitespace
    if not context_texts.strip():
        return "Error: Context is empty or improperly formatted.", None
    
    # Split the long context text into chunks using the chunking function
    context_chunks = get_chunks(context_texts, chunk_size)
    
    # Handle single context case
    if len(context_chunks) == 1:
        return context_chunks[0], 1.0  # Return the single chunk with perfect similarity
    
    # Encode the user input to create a query embedding
    user_embedding = encode_text(user_input).reshape(1, -1)
    
    # Encode all context chunks to create embeddings
    chunk_embeddings = np.array([encode_text(chunk) for chunk in context_chunks])
    
    # Compute cosine similarity between the user input and each chunk
    similarities = cosine_similarity(user_embedding, chunk_embeddings).flatten()
    
    # Check if any similarity scores are above the threshold
    if max(similarities) < similarity_threshold:
        return "No relevant context found for the user input.", None
    
    # Identify the most relevant chunk based on the highest cosine similarity score
    most_relevant_idx = np.argmax(similarities)
    most_relevant_chunk = context_chunks[most_relevant_idx]
    
    # Return the most relevant chunk and the similarity score
    return most_relevant_chunk


def extract_keywords_with_gpt(user_input, max_tokens=100, temperature=0.3):
    # Define a prompt to ask GPT-4 to extract keywords and important terms
    keyword_prompt = f"Extract the most important keywords, scientific concepts, and parameters from the following user query:\n\n{user_input}"
    
    # Call GPT-4 to extract keywords based on the user prompt
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in identifying key scientific terms and concepts."},
            {"role": "user", "content": keyword_prompt}
        ],
        max_tokens=max_tokens,
        temperature=temperature
    )
    
    # Extract the content from GPT-4's reply
    extracted_keywords = response.choices[0].message.content.strip()
    
    return extracted_keywords

def fetch_nasa_ads_references(prompt):
    try:
        # Use the entire prompt for the query
        simplified_query = prompt

        # Query NASA ADS for relevant papers
        papers = ADS.query_simple(simplified_query)
        
        if not papers or len(papers) == 0:
            return [("No results found", "N/A", "N/A")]
        
        # Include authors in the references
        references = [
            (
                paper['title'][0], 
                ", ".join(paper['author'][:3]) + (" et al." if len(paper['author']) > 3 else ""), 
                paper['bibcode']
            ) 
            for paper in papers[:5]  # Limit to 5 references
        ]
        return references
    
    except Exception as e:
        return [("Error fetching references", str(e), "N/A")]

def fetch_exoplanet_data():
    # Connect to NASA Exoplanet Archive TAP Service
    tap_service = vo.dal.TAPService("https://exoplanetarchive.ipac.caltech.edu/TAP")

    # Query to fetch all columns from the pscomppars table
    ex_query = """
        SELECT TOP 10 pl_name, hostname, sy_snum, sy_pnum, discoverymethod, disc_year, disc_facility, pl_controv_flag, pl_orbper, pl_orbsmax, pl_rade, pl_bmasse, pl_orbeccen, pl_eqt, st_spectype, st_teff, st_rad, st_mass, ra, dec, sy_vmag
        FROM pscomppars
    """
    # Execute the query
    qresult = tap_service.search(ex_query)

    # Convert to a Pandas DataFrame
    ptable = qresult.to_table()
    exoplanet_data = ptable.to_pandas()

    return exoplanet_data

def generate_response(user_input, relevant_context="", references=[], max_tokens=150, temperature=0.7, top_p=0.9, frequency_penalty=0.5, presence_penalty=0.0):
    if relevant_context:
        combined_input = f"Scientific Context: {relevant_context}\nUser Input: {user_input}\nPlease generate a detailed structured response as per the defined sections and table format."
    else:
        combined_input = f"User Input: {user_input}\nPlease generate a detailed structured response as per the defined sections and table format."
    
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": combined_input}
        ],
        max_tokens=max_tokens,
        temperature=temperature,
        top_p=top_p,
        frequency_penalty=frequency_penalty,
        presence_penalty=presence_penalty
    )
    
    # Append references to the response
    if references:
        response_content = response.choices[0].message.content.strip()
        references_text = "\n\nADS References:\n" + "\n".join(
            [f"- {title} by {authors} (Bibcode: {bibcode})" for title, authors, bibcode in references]
        )
        return f"{response_content}\n{references_text}"
    
    return response.choices[0].message.content.strip()

def generate_data_insights(user_input, exoplanet_data, max_tokens=500, temperature=0.3):
    """
    Generate insights by passing the user's input along with the exoplanet data to GPT-4.
    """
    # Convert the dataframe to a readable format for GPT (e.g., CSV-style text)
    data_as_text = exoplanet_data.to_csv(index=False)  # CSV-style for better readability

    # Create a prompt with the user query and the data sample
    insights_prompt = (
        f"Analyze the following user query and provide relevant insights based on the provided exoplanet data.\n\n"
        f"User Query: {user_input}\n\n"
        f"Exoplanet Data:\n{data_as_text}\n\n"
        f"Please provide insights that are relevant to the user's query."
    )
    
    # Call GPT-4 to generate insights based on the data and user input
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in analyzing astronomical data and generating insights."},
            {"role": "user", "content": insights_prompt}
        ],
        max_tokens=max_tokens,
        temperature=temperature
    )
    
    # Extract and return GPT-4's insights
    data_insights = response.choices[0].message.content.strip()
    return data_insights


def export_to_word(response_content, subdomain_definition, science_goal):
    doc = Document()
    
    # Add a title (optional, you can remove this if not needed)
    doc.add_heading('AI Generated SCDD', 0)

    # Insert the Subdomain Definition at the top
    doc.add_heading('Subdomain Definition:', level=1)
    doc.add_paragraph(subdomain_definition)

    # Insert the Science Goal at the top
    doc.add_heading('Science Goal:', level=1)
    doc.add_paragraph(science_goal)

    # Split the response into sections based on ### headings
    sections = response_content.split('### ')
    
    for section in sections:
        if section.strip():
            # Handle the "Technical Requirements Table" separately with proper formatting
            if section.startswith('Technical Requirements Table'):
                doc.add_heading('Technical Requirements Table', level=1)
                
                # Extract table lines
                table_lines = section.split('\n')[2:]  # Start after the heading line
                
                # Check if it's an actual table (split lines by '|' symbol)
                table_data = [line.split('|')[1:-1] for line in table_lines if '|' in line]
                
                if table_data:
                    # Add table to the document
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            cell = table.cell(i, j)
                            cell.text = cell_text.strip()
                            # Apply text wrapping for each cell
                            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcW w:w="2500" w:type="pct" ' + nsdecls('w') + '/>'))
                
                # Process any paragraphs that follow the table
                paragraph_after_table = '\n'.join([line for line in table_lines if '|' not in line and line.strip()])
                if paragraph_after_table:
                    doc.add_paragraph(paragraph_after_table.strip())
            
            # Handle the "ADS References" section
            elif section.startswith('ADS References'):
                doc.add_heading('ADS References', level=1)
                references = section.split('\n')[1:]  # Skip the heading
                for reference in references:
                    if reference.strip():
                        doc.add_paragraph(reference.strip())
            
            # Add all other sections as plain paragraphs
            else:
                doc.add_paragraph(section.strip())
    
    # Save the document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    
    return temp_file.name

def extract_table_from_response(gpt_response):
    # Split the response into lines
    lines = gpt_response.strip().split("\n")
    
    # Find where the table starts and ends (based on the presence of pipes `|` and at least 3 columns)
    table_lines = [line for line in lines if '|' in line and len(line.split('|')) > 3]
    
    # If no table is found, return None or an empty string
    if not table_lines:
        return None
    
    # Find the first and last index of the table lines
    first_table_index = lines.index(table_lines[0])
    last_table_index = lines.index(table_lines[-1])
    
    # Extract only the table part
    table_text = lines[first_table_index:last_table_index + 1]
    
    return table_text

def gpt_response_to_dataframe(gpt_response):
    # Extract the table text from the GPT response
    table_lines = extract_table_from_response(gpt_response)
    
    # If no table found, return an empty DataFrame
    if table_lines is None or len(table_lines) == 0:
        return pd.DataFrame()

    # Find the header and row separator (assume it's a line with dashes like |---|)
    try:
        # The separator line (contains dashes separating headers and rows)
        sep_line_index = next(i for i, line in enumerate(table_lines) if set(line.strip()) == {'|', '-'})
    except StopIteration:
        # If no separator line is found, return an empty DataFrame
        return pd.DataFrame()

    # Extract headers (the line before the separator) and rows (lines after the separator)
    headers = [h.strip() for h in table_lines[sep_line_index - 1].split('|')[1:-1]]
    
    # Extract rows (each line after the separator)
    rows = [
        [cell.strip() for cell in row.split('|')[1:-1]]
        for row in table_lines[sep_line_index + 1:]
    ]

    # Create DataFrame
    df = pd.DataFrame(rows, columns=headers)
    return df
    
def chatbot(user_input, context="", subdomain="", use_encoder=False, max_tokens=150, temperature=0.7, top_p=0.9, frequency_penalty=0.5, presence_penalty=0.0):
    if use_encoder and context:
        context_texts = context
        relevant_context = retrieve_relevant_context(user_input, context_texts)
    else:
        relevant_context = ""

    # Fetch NASA ADS references using the full prompt
    references = fetch_nasa_ads_references(subdomain)

    # Generate response from GPT-4
    response = generate_response(user_input, relevant_context, references, max_tokens, temperature, top_p, frequency_penalty, presence_penalty)

    # Export the response to a Word document
    word_doc_path = export_to_word(response, subdomain, user_input)

    # Fetch exoplanet data
    exoplanet_data = fetch_exoplanet_data()

    # Generate insights based on the user query and exoplanet data
    data_insights = generate_data_insights(user_input, exoplanet_data)

    # Extract and convert the table from the GPT-4 response into a DataFrame
    extracted_table_df = gpt_response_to_dataframe(response)

    # Combine the response and the data insights
    full_response = f"{response}\n\nInsights from Existing Data: {data_insights}"
    
    # Embed Miro iframe
    iframe_html = """
    <iframe width="768" height="432" src="https://miro.com/app/live-embed/uXjVKuVTcF8=/?moveToViewport=-331,-462,5434,3063&embedId=710273023721" frameborder="0" scrolling="no" allow="fullscreen; clipboard-read; clipboard-write" allowfullscreen></iframe>
    """
    
    mapify_button_html = """
    <style>
        .mapify-button {
            background: linear-gradient(135deg, #1E90FF 0%, #87CEFA 100%);
            border: none;
            color: white;
            padding: 15px 35px;
            text-align: center;
            text-decoration: none;
            display: inline-block;
            font-size: 18px;
            font-weight: bold;
            margin: 20px 2px;
            cursor: pointer;
            border-radius: 25px;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2);
        }
        .mapify-button:hover {
            background: linear-gradient(135deg, #4682B4 0%, #1E90FF 100%);
            box-shadow: 0 6px 20px rgba(0, 0, 0, 0.3);
            transform: scale(1.05);
        }
    </style>
    <a href="https://mapify.so/app/new" target="_blank">
        <button class="mapify-button">Create Mind Map on Mapify</button>
    </a>
    """
    return full_response, extracted_table_df, word_doc_path, exoplanet_data, iframe_html, mapify_button_html

iface = gr.Interface(
    fn=chatbot,
    inputs=[
        gr.Textbox(lines=5, placeholder="Enter your Science Goal...", label="Science Goal"),
        gr.Textbox(lines=10, placeholder="Enter Context Text...", label="Context"),
        gr.Textbox(lines=2, placeholder="Define your Subdomain...", label="Subdomain Definition"),
        gr.Checkbox(label="Use NASA SMD Bi-Encoder for Context"),
        gr.Slider(50, 2000, value=150, step=10, label="Max Tokens"),
        gr.Slider(0.0, 1.0, value=0.7, step=0.1, label="Temperature"),
        gr.Slider(0.0, 1.0, value=0.9, step=0.1, label="Top-p"),
        gr.Slider(0.0, 1.0, value=0.5, step=0.1, label="Frequency Penalty"),
        gr.Slider(0.0, 1.0, value=0.0, step=0.1, label="Presence Penalty")
    ],
    outputs=[
        gr.Textbox(label="ExosAI finds..."),                           
        gr.Dataframe(label="SC Requirements Table"),      
        gr.File(label="Download SCDD", type="filepath"),                
        gr.Dataframe(label="Exoplanet Data Table"),                     
        gr.HTML(label="Miro"),                                          
        gr.HTML(label="Generate Mind Map on Mapify") 
    ],
    title="ExosAI - NASA SMD SCDD AI Assistant [version-0.9a]",
    description="ExosAI is an AI-powered assistant for generating and visualising HWO Science Cases",
)

iface.launch(share=True)